import streamlit as st
import io
from PIL import Image
import base64
import os
import pandas as pd
import datetime
import requests
import re
import time
import yaml
import bcrypt
import duckdb
import json
import traceback
from langchain_groq import ChatGroq
from langchain.tools import tool
from functions import gen_sql,SQL_AGENT_SYSTEM_PROMPT,get_db_metadata,validate_sql,run_query_on_duckdb,gen_nl

st.set_page_config(layout="wide")

def load_config():
    with open('config.yaml') as file:
        return yaml.safe_load(file)

def verify_password(stored_hash, password):
    return bcrypt.checkpw(password.encode('utf-8'), stored_hash.encode('utf-8'))

# ===== LANGCHAIN AGENT FOR CSV/EXCEL Q&A =====
def create_csv_excel_qa_agent(df, file_type="csv"):
    """Create a data analyzer for CSV/Excel Q&A using LLM with context"""
    
    def analyze_with_context(question: str) -> str:
        """Analyze the dataframe and answer user questions"""
        try:
            llm = ChatGroq(
                model="lama-3.1-8b-instant",#"llama-3.3-70b-versatile",
                api_key="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",
                temperature=0.7,
                max_tokens=2048,
            )
            
            # Prepare data context
            df_info = f"""Dataset Information:
- Shape: {df.shape[0]} rows, {df.shape[1]} columns
- Columns: {list(df.columns)}
- Data types: {dict(df.dtypes)}
- Column descriptions:
"""
            for col in df.columns:
                dtype = df[col].dtype
                if dtype in ['float64', 'int64']:
                    df_info += f"  - {col} ({dtype}): min={df[col].min()}, max={df[col].max()}, mean={df[col].mean():.2f}\n"
                else:
                    df_info += f"  - {col} ({dtype}): {df[col].nunique()} unique values\n"
            
            # Add sample data
            df_info += f"\nFirst 5 rows:\n{df.head().to_string()}\n"
            
            # Build analysis prompt
            analysis_prompt = f"""You are an expert data analyst. Your task is to analyze the dataset and answer the user's question.

{df_info}

User Question: {question}

Instructions:
1. Analyze the data thoroughly - perform aggregations, transformations, and calculations as needed to find the answer
2. Use statistics, filtering, grouping, and other data operations to extract insights
3. Show your analytical reasoning and the data operations you performed
4. Once you have analyzed and determined the answer, provide ONLY a simple, natural language response

Format your response as follows:
---ANALYSIS---
[Show your data analysis, aggregations, transformations, calculations, and reasoning here]

---FINAL ANSWER---
[Provide ONLY a simple, natural language response to the question. No code, no technical details, just the answer.]

The final answer should be:
- Direct and concise (2-5 sentences)
- In plain language that a non-technical person can understand
- Include relevant numbers/metrics if applicable
- Address the user's question completely"""
            
            response = llm.invoke(analysis_prompt)
            full_response = response.content
            
            # Extract only the final answer section
            if "---FINAL ANSWER---" in full_response:
                final_answer = full_response.split("---FINAL ANSWER---")[-1].strip()
                return final_answer
            else:
                return full_response
            
        except Exception as e:
            return f"Error analyzing data: {str(e)}"
    
    return analyze_with_context

config = load_config()

users = config['credentials']['usernames']


if 'logged_in' not in st.session_state or not st.session_state.logged_in:
    print('Yessss')
    col1, col2, col3 = st.columns([1, 2, 1])  # Middle column is wider

    with col2:  # Place the form in the center
        st.title(" ")

        username_input = st.text_input("Username")
        
        password_input = st.text_input("Password", type="password")

        if st.button('Login'):
            if username_input in users:
                user = users[username_input]
                stored_password_hash = user['password']

                
                if verify_password(stored_password_hash, password_input):
                    st.session_state.logged_in = True
                    st.session_state.username = user
                    
                    st.success(f"Welcome {user['name']}!")
                    st.rerun()
                else:
                    st.error('Invalid password.')
            else:
                st.error('Username not found.')

else:
    user = st.session_state.username
    st.session_state.username1 = user['name']

    st.markdown("""
    <style>
        .stDownloadButton>button {
            font-size: 12px;  /* Adjust font size */
            padding: 4px 8px;  /* Adjust padding */
            height: auto;
            display: flex;
            align-items: center;
            justify-content: center;
            cursor: pointer;
        }

        .stDownloadButton {
            position: fixed;
            top: 10px;     /* Distance from the top */
            right: 110px;   /* Distance from the right */
            z-index: 999;  /* Ensure the button is above other elements */
        }
    </style>
""", unsafe_allow_html=True)


st.markdown("""
        <style>
               .block-container {
                    padding-top: 0rem;
                    padding-bottom: 0rem;
                }
        </style>
        """, unsafe_allow_html=True)



st.markdown("""
<style>


[data-baseweb="base-input"]{
background:lightgray;
border: 2px;
border-radius: 3px;
}

input[class]{
font-weight: bold;
font-size:120%;
color: black;
}
</style>
""", unsafe_allow_html=True)



if 'logged_in' in st.session_state :

    if "prompts" not in st.session_state:
        st.session_state["prompts"] = []
        st.session_state['initial_prompt'] =[]
    if "answers" not in st.session_state:
        st.session_state["answers"] = []
        
    if "insight" not in st.session_state:
        st.session_state["insight"] = []

    st.session_state["chat_hist"] = {}

    if "fin_sql_query" not in st.session_state:
        st.session_state['fin_sql_query'] = []
    if "fin_msk_df" not in st.session_state:
        st.session_state['fin_msk_df'] = []
    
    # Session state for Summarized Narrative Generator
    if "summarizer_file" not in st.session_state:
        st.session_state["summarizer_file"] = None
    if "summarizer_content" not in st.session_state:
        st.session_state["summarizer_content"] = ""
    if "summarizer_result" not in st.session_state:
        st.session_state["summarizer_result"] = ""
    
    def new_chat():
        st.session_state["prompts"] = []
        st.session_state["chat_hist"] = {}
        st.session_state["answers"] = []
        st.session_state["insight"] = []
        
        st.session_state.query = ''
        st.session_state['initial_prompt'] =[]

        st.session_state["fin_sql_query"] = []
        st.session_state["fin_msk_df"] = []

        


    col1, col2, col3 = st.columns([4,2,2])

    st.markdown(
        """
        <style>
        .block-container {
            margin-top: 0px !important;
            padding-top: 10px !important;
        }

        .stButton {
            margin: 0 !important;
            padding: 0 !important;
            display: flex !important;
            justify-content: flex-end !important;
        }

        .stButton>button {
            padding: 10px 20px !important;
            font-size: 14px !important;
            height: 40px !important;
            width: auto !important;
            margin: 0 !important;
            background-color: #1f77b4 !important;
            color: white !important;
            border-radius: 5px !important;
            border: 1px solid #0d47a1 !important;
            font-weight: bold !important;
            cursor: pointer !important;
        }

        .stButton>button:hover {
            background-color: #0d47a1 !important;
            transform: scale(1.05) !important;
        }
        """, unsafe_allow_html=True
    )

   

    with col1:
        st.subheader('Insight Generator')  # Set image width

    
    with col3:
        st.button("New Chat", on_click=new_chat, type='primary')
    
    # Create tabs for two separate features
    tab1, tab2 = st.tabs(["📊 Insight Q&A", "📄 Summarized Narrative Generator"])
    
    placeholder_fdbk = st.columns(1)
    st.session_state.plc_fdbk = placeholder_fdbk[0].empty()

   
    
    # ===== TAB 1: INSIGHT Q&A =====
    with tab1:
        # Create containers specific to Tab 1
        st.session_state.col2_dld = st.empty()
        st.session_state.response_container = st.container()
        st.session_state.info_placeholder = st.empty()
        st.session_state.response_exp_sql = st.empty()
        st.session_state.response_exp_dta = st.empty()
        st.session_state.response_exp_res = st.empty()
        
        
        
        
        # Chat input at bottom
        st.session_state.query =  st.chat_input("Input your question here...")
        
        if st.session_state.query != None:

            prompt = st.session_state.query
            
            initial_prompt = prompt   
            initial_prompt_check = prompt

            reframe_prompt = prompt
            DB_PATH = "salesdata.duckdb"
            
            llm = ChatGroq(
                    model="llama-3.3-70b-versatile",      # e.g. "llama3-70b-8192" or 
                    api_key="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",  # your Groq API key
                    temperature=0,
                    max_tokens=None,
                    timeout=None,
                    max_retries=2,
                )
            with st.spinner("Thinking ... ⏳"):
                st.session_state.info_placeholder.info(initial_prompt)
                
                if len(prompt)>0:
                    try:
                    
                        # st.session_state['tables_list'], st.session_state['columns_map'], st.session_state['sample_rows_map'] = get_db_metadata(DB_PATH,5)
                        tables_list, columns_map, sample_rows_map = get_db_metadata(DB_PATH,5)
                        sql_query = gen_sql(reframe_prompt, tables_list, columns_map, sample_rows_map, SQL_AGENT_SYSTEM_PROMPT, llm)
                        sql_query = validate_sql(sql_query, columns_map, llm)
                        # st.write("Generated SQL Query: ")
                        # st.code(sql_query)
                    except Exception as e: 
                        print('error',e)
                        
                        st.session_state.nl_response = "Oops! error while generating SQL query. Please try again."
                        
                        print(e)
                        print('.......DAX..continue...')
                        
                        st.session_state['initial_prompt'].append(initial_prompt)    
                        st.session_state["prompts"].append(reframe_prompt)                               
                        st.session_state["answers"].append(st.session_state.nl_response)
                        
                        st.session_state["fin_sql_query"].append('OOPS! No query generated')
                        st.session_state["fin_msk_df"].append('Error,Please try again')
                        st.session_state.zipped_data = list(zip(st.session_state['initial_prompt'],st.session_state["prompts"],st.session_state["answers"],st.session_state["fin_sql_query"],st.session_state["fin_msk_df"]))
                           
                    else:
                        st.session_state.final_query = sql_query
                        
                        if len(st.session_state.final_query) > 0:
                            try:
                                
                                df = run_query_on_duckdb(sql_query, db_path=DB_PATH)
                                df_json = df.to_json(orient='records', lines=True)
                                # st.table(df.head(5))
                            except Exception as e: 
                                print("came in sql exe except")
                                
                                print(e)
                                st.session_state.nl_response = "Oops! error while executing SQL query. Please try again."
                            
                                st.session_state['initial_prompt'].append(initial_prompt)    
                                st.session_state["prompts"].append(reframe_prompt)                               
                                st.session_state["answers"].append(st.session_state.nl_response)
                              
                                st.session_state["fin_sql_query"].append(st.session_state.final_query)
                                st.session_state["fin_msk_df"].append('Error,Please try again')
                                st.session_state.zipped_data = list(zip(st.session_state['initial_prompt'],st.session_state["prompts"],st.session_state["answers"],st.session_state["fin_sql_query"],st.session_state["fin_msk_df"]))
                           
                            else:
                                if df.isna().all().all():
                                    print("came in all null check block")
                              
                        
                                    st.session_state.nl_response = "Oops! there is no data to provide response for query."
                                    st.session_state['initial_prompt'].append(initial_prompt)    
                                    st.session_state["prompts"].append(reframe_prompt)                               
                                    st.session_state["answers"].append(st.session_state.nl_response)
                                    
                                    st.session_state["fin_sql_query"].append(st.session_state.final_query)
                                    df_dic = df.to_dict(orient='records') 
                                    st.session_state["fin_msk_df"].append(df_dic)
                                    st.session_state.zipped_data = list(zip(st.session_state['initial_prompt'],st.session_state["prompts"],st.session_state["answers"],st.session_state["fin_sql_query"],st.session_state["fin_msk_df"]))
                           
                                else:
                                
                                    try:
                                      
                                        nlreponse = gen_nl(reframe_prompt, df_json, llm)
                                        # st.write("Generated Response: ",nlreponse)
                                        
                                        # Append to session state on successful response
                                        st.session_state['initial_prompt'].append(initial_prompt)
                                        st.session_state["prompts"].append(reframe_prompt)
                                        st.session_state["answers"].append(nlreponse)
                                        st.session_state["fin_sql_query"].append(st.session_state.final_query)
                                        df_dic = df.to_dict(orient='records')
                                        st.session_state["fin_msk_df"].append(df_dic)
                                        st.session_state.zipped_data = list(zip(st.session_state['initial_prompt'],st.session_state["prompts"],st.session_state["answers"],st.session_state["fin_sql_query"],st.session_state["fin_msk_df"]))

                                    except Exception as e: 
                                        st.session_state.nl_response = 'Oops! It is taking longer than expected. Please ask the question again or try asking another question.'
                                        
                                        print("Error: ", e)
                                        st.warning("It is taking longer than expected. Please ask the question again or try asking another question.")
                                        st.session_state['initial_prompt'].append(initial_prompt)  
                                        st.session_state["prompts"].append(reframe_prompt)                               
                                        st.session_state["answers"].append(st.session_state.nl_response)
                                        
                                        st.session_state["fin_sql_query"].append(st.session_state.final_query)
                                        df_dic = df.to_dict(orient='records') 
                                        st.session_state["fin_msk_df"].append(df_dic)
                                        st.session_state.zipped_data = list(zip(st.session_state['initial_prompt'],st.session_state["prompts"],st.session_state["answers"],st.session_state["fin_sql_query"],st.session_state["fin_msk_df"]))
        print("st.session_state['prompts'] : ",st.session_state['prompts'])
        if st.session_state['prompts']:
            print('came in prompt if')
            with st.session_state.response_container:
                for index, (initial_q, prompt, response, query, data) in enumerate((st.session_state.zipped_data)):

                    st.session_state.info_placeholder.empty()
                    st.session_state.response_exp_sql.empty()
                    st.session_state.response_exp_dta.empty()
                    st.session_state.response_exp_res.empty()
                    response = response.replace("$","\\$")
                    
                    # Use emoji icons instead of image files
                    user_emoji = "👤"  # User icon
                    assistant_emoji = "🤖"  # AI/Assistant icon
                    
                    col1, col2 = st.columns([15,1])

                    with col1:
                        st.markdown(
                            f"""
                            <div style="float: right; clear: both; margin-top: 15px; padding: 10px; background-color: #f0f0f0; 
                            border: 1px solid #ddd; border-radius: 10px; display: inline-block; margin-bottom: 20px;
                            word-wrap: break-word;">
                                {initial_q}
                            </div>
                            """,
                            unsafe_allow_html=True
                        )

                    with col2:
                        st.write(user_emoji)

                    
                    with st.expander("View Details"):
                        st.code(query)
                        # Handle DataFrame construction safely
                        if isinstance(data, list) and len(data) > 0:
                            df_display = pd.DataFrame(data)
                            st.dataframe(df_display)

                        else:
                            st.warning("No data available to display")

                    st.chat_message(name = "assistant", avatar="🤖").markdown(response)                       
    
    # ===== TAB 2: SUMMARIZED NARRATIVE GENERATOR (with RAG for documents) =====
    with tab2:
        st.markdown("""
        ### 📄 Summarized Narrative Generator
        Upload any file (CSV, XLSX, TXT, PDF, DOC, DOCX, JSON) to generate a concise narrative summary and insights.
        - **CSV/XLSX**: Direct summarization
        - **Documents (TXT, PDF, DOC, DOCX)**: RAG-based summarization for better context understanding
        """)
        
        uploaded_file = st.file_uploader("Choose a file", type=["csv", "txt", "pdf", "xlsx", "json", "doc", "docx"])
        
        if uploaded_file is not None:
            st.session_state["summarizer_file"] = uploaded_file.name
            file_type = uploaded_file.type
            
            # Read file content based on type
            try:
                if file_type == "text/plain":
                    st.session_state["summarizer_content"] = uploaded_file.read().decode("utf-8")
                    is_document = True
                elif file_type == "text/csv":
                    df_temp = pd.read_csv(uploaded_file)
                    st.session_state["summarizer_content"] = df_temp.to_string()
                    is_document = False
                elif file_type == "application/json":
                    json_data = json.load(uploaded_file)
                    st.session_state["summarizer_content"] = json.dumps(json_data, indent=2)
                    is_document = False
                elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                    df_temp = pd.read_excel(uploaded_file)
                    st.session_state["summarizer_content"] = df_temp.to_string()
                    is_document = False
                elif file_type == "application/pdf":
                    try:
                        import PyPDF2
                        pdf_reader = PyPDF2.PdfReader(uploaded_file)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                        st.session_state["summarizer_content"] = text
                        is_document = True
                    except ImportError:
                        st.warning("PDF support requires PyPDF2. Install it using: pip install PyPDF2")
                        st.session_state["summarizer_content"] = ""
                        is_document = False
                elif file_type == "application/msword" or file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    # Handle DOC and DOCX files
                    try:
                        from docx import Document
                        
                        if file_type == "application/msword":
                            # For older .doc files, try using python-docx or fallback
                            st.warning("DOC format detected. For best results, please convert to DOCX format.")
                            try:
                                doc = Document(uploaded_file)
                                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                                st.session_state["summarizer_content"] = text
                                is_document = True
                            except:
                                st.error("Unable to process .doc file. Please convert to .docx format.")
                                st.session_state["summarizer_content"] = ""
                                is_document = False
                        else:
                            # DOCX file
                            doc = Document(uploaded_file)
                            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            st.session_state["summarizer_content"] = text
                            is_document = True
                    except ImportError:
                        st.warning("Word document support requires python-docx. Install it using: pip install python-docx")
                        st.session_state["summarizer_content"] = ""
                        is_document = False
                else:
                    is_document = False
                
                if st.session_state["summarizer_content"]:
                    st.success(f"✅ File loaded successfully: {uploaded_file.name}")
                    
                    # Preview section
                    with st.expander("Preview File Content"):
                        preview_length = min(500, len(st.session_state["summarizer_content"]))
                        st.text_area("Content Preview", st.session_state["summarizer_content"][:preview_length], height=150, disabled=True)
                    
                    # Generate summary button
                    col_generate, col_clear = st.columns([1, 1])
                    
                    with col_generate:
                        if st.button("🔍 Generate Summary & Insights", key="generate_summary"):
                            with st.spinner("Generating narrative summary... 📝"):
                                try:
                                    llm_summarizer = ChatGroq(
                                        model="llama-3.3-70b-versatile",
                                        api_key="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",
                                        temperature=0.7,
                                        max_tokens=2048,
                                        timeout=None,
                                        max_retries=2,
                                    )
                                    
                                    if is_document:
                                        # RAG approach for documents - using simple chunking
                                        try:
                                            # Split document into chunks manually to avoid dependency conflicts
                                            content = st.session_state["summarizer_content"]
                                            
                                            # Split by paragraphs or sentences
                                            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                                            
                                            # Create chunks (limit size)
                                            chunks = []
                                            current_chunk = ""
                                            for para in paragraphs:
                                                if len(current_chunk) + len(para) < 2000:
                                                    current_chunk += para + "\n\n"
                                                else:
                                                    if current_chunk:
                                                        chunks.append(current_chunk)
                                                    current_chunk = para + "\n\n"
                                            if current_chunk:
                                                chunks.append(current_chunk)
                                            
                                            # Use top chunks (instead of embedding-based retrieval)
                                            retrieved_context = "\n\n".join(chunks[:5])  # Top 5 chunks
                                            
                                            summary_prompt = f"""You are a senior business analyst and subject matter expert. Analyze the following document content in depth and provide a comprehensive, detailed narrative that goes beyond surface-level observations.

                                                            Provide the following sections in your analysis:

                                                            1. **Executive Summary** (4-5 sentences)
                                                            - Provide a clear, concise overview of the document's main purpose, scope, and key takeaways
                                                            - Highlight the most critical information that stakeholders should know

                                                            2. **Key Insights & Strategic Findings** (8-10 bullet points)
                                                            - Identify the most important insights, patterns, and trends
                                                            - Explain the significance and implications of each finding
                                                            - Connect insights to broader business or contextual implications
                                                            - Highlight any critical metrics, percentages, or quantitative data

                                                            3. **Core Themes & Topics** (6-8 points)
                                                            - Identify and elaborate on the main themes and recurring topics
                                                            - Explain how these themes interconnect and relate to each other
                                                            - Provide context for understanding the document's focus areas

                                                            4. **Risk Factors & Challenges** (if applicable)
                                                            - Identify any potential risks, challenges, or areas of concern mentioned or implied
                                                            - Explain the impact and severity of each identified issue

                                                            5. **Opportunities & Recommendations** (6-8 actionable points)
                                                            - Suggest specific, actionable recommendations based on the document's content
                                                            - Provide concrete next steps and priorities
                                                            - Explain the expected benefits and outcomes of implementing recommendations

                                                            6. **Supporting Data & Evidence**
                                                            - Highlight important statistics, metrics, or evidence that support the key findings
                                                            - Include any notable quotes or specific examples from the document

                                                            Document Content:
                                                            {retrieved_context}

                                                        Guidelines:
                                                        - Be specific and detailed in your analysis
                                                        - Use clear, professional language
                                                        - Organize information hierarchically from most to least important
                                                        - Ensure all recommendations are actionable and specific
                                                        - Format with clear headings, subheadings, and bullet points for easy reading
                                                        - Provide meaningful context and explanation for each point
                                                        - Connect individual findings to larger strategic implications"""
                                        
                                        except Exception as e:
                                            # Fallback to simple summarization
                                            st.warning(f"Using simplified summarization: {str(e)}")
                                            summary_prompt = f"""You are an expert analyst. Analyze the following content and provide:
                                                1. A concise executive summary (3-4 sentences)
                                                2. Key insights and findings (bullet points)
                                                3. Important metrics or statistics
                                                4. Recommendations (if applicable)

                                                Content to analyze:
                                                {st.session_state['summarizer_content'][:3000]}

                                                Please format your response clearly with proper headings and bullet points."""

                                        #             Content to analyze:
                                        #             {st.session_state['summarizer_content'][:3000]}

                                        #             Please format your response clearly with proper headings and bullet points."""
                                    else:
                                        # Direct summarization for CSV/XLSX/JSON
                                        summary_prompt = f"""You are an expert analyst. Analyze the following content and provide:
                                            1. A concise executive summary (3-4 sentences)
                                            2. Key insights and findings (bullet points)
                                            3. Important metrics or statistics
                                            4. Recommendations (if applicable)

                                            Content to analyze:
                                            {st.session_state['summarizer_content'][:3000]}

                                            Please format your response clearly with proper headings and bullet points."""
                                    
                                    response = llm_summarizer.invoke(summary_prompt)
                                    st.session_state["summarizer_result"] = response.content
                                    
                                except Exception as e:
                                    st.error(f"Error generating summary: {str(e)}")
                                    print(f"Summarizer Error: {e}")
                    
                    with col_clear:
                        if st.button("🗑️ Clear All", key="clear_summarizer"):
                            st.session_state["summarizer_file"] = None
                            st.session_state["summarizer_content"] = ""
                            st.session_state["summarizer_result"] = ""
                            st.session_state["tab2_qa_history"] = []
                            st.session_state["summarizer_dataframe"] = None
                            st.rerun()
                    
                    # Display results
                    if st.session_state["summarizer_result"]:
                        st.markdown("---")
                        st.markdown("### 📋 Generated Summary & Insights")
                        
                        # Create tabs for different views
                        summary_tab1, summary_tab2 = st.tabs(["Summary", "Raw Output"])
                        
                        with summary_tab1:
                            st.markdown(st.session_state["summarizer_result"])
                        
                        with summary_tab2:
                            st.text_area("Raw Output", st.session_state["summarizer_result"], height=300, disabled=True)
                        
                        # Download summary
                        summary_text = st.session_state["summarizer_result"]
                        st.download_button(
                            label="⬇️ Download Summary as Text",
                            data=summary_text,
                            file_name=f"summary_{uploaded_file.name}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain"
                        )
                    
                    # ===== Q&A SECTION =====
                    st.markdown("---")
                    st.markdown("### ❓ Ask Questions About This Document")
                    
                    # Initialize Q&A history for this document
                    if "tab2_qa_history" not in st.session_state:
                        st.session_state["tab2_qa_history"] = []
                    
                    # Create a container for Q&A history that won't shift around
                    qa_history_container = st.container()
                    
                    # Question input stays in fixed position
                    qa_question = st.chat_input("Ask a question about the document...", key="tab2_qa_input")
                    
                    # Display Q&A history in the container
                    with qa_history_container:
                        if st.session_state["tab2_qa_history"]:
                            st.markdown("#### 💬 Q&A History")
                            for qa_pair in (st.session_state["tab2_qa_history"]):
                                col1, col2 = st.columns([15, 1])
                                with col1:
                                    st.markdown(
                                        f"""
                                        <div style="padding: 10px; background-color: #e3f2fd; 
                                        border: 1px solid #90caf9; border-radius: 10px; margin-bottom: 10px;
                                        word-wrap: break-word;">
                                            <strong>Q:</strong> {qa_pair['question']}
                                        </div>
                                        """,
                                        unsafe_allow_html=True
                                    )
                                with col2:
                                    st.write("❓")
                                
                                st.chat_message(name="assistant", avatar="🤖").markdown(qa_pair['answer'])
                    
                    if qa_question:
                        with st.spinner("Analyzing document and generating answer... 🔍"):
                            try:
                                # Check if file is CSV or Excel for agent-based Q&A
                                file_ext = uploaded_file.name.lower().split('.')[-1]
                                
                                if file_ext in ['csv', 'xlsx', 'xls']:
                                    # Use LangChain Agent for CSV/Excel
                                    try:
                                        # Load dataframe if not already loaded
                                        if "summarizer_dataframe" not in st.session_state or st.session_state["summarizer_dataframe"] is None:
                                            if file_ext == 'csv':
                                                # Seek to beginning of file before reading
                                                uploaded_file.seek(0)
                                                st.session_state["summarizer_dataframe"] = pd.read_csv(uploaded_file)
                                            else:
                                                # Seek to beginning of file before reading
                                                uploaded_file.seek(0)
                                                st.session_state["summarizer_dataframe"] = pd.read_excel(uploaded_file)
                                        
                                        df = st.session_state["summarizer_dataframe"]
                                        
                                        # Verify dataframe is not None
                                        if df is None:
                                            st.error("Failed to load dataframe. Please try uploading the file again.")
                                            raise ValueError("Dataframe is None after loading")
                                        
                                        # Create and execute agent
                                        analyzer = create_csv_excel_qa_agent(df, file_type=file_ext)
                                        
                                        if analyzer:
                                            # Run analyzer with the question
                                            answer = analyzer(qa_question)
                                        else:
                                            # Fallback to direct LLM if agent creation fails
                                            llm_qa = ChatGroq(
                                                model="llama-3.3-70b-versatile",
                                                api_key="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",
                                                temperature=0.7,
                                                max_tokens=1500,
                                            )
                                            qa_prompt = f"""You are a data analyst. The user has a question about their data file.
Question: {qa_question}

Data Summary:
- Shape: {df.shape[0]} rows, {df.shape[1]} columns
- Columns: {list(df.columns)}
- Data types: {dict(df.dtypes)}

First few rows:
{df.head(3).to_string()}

Please provide a helpful, detailed answer based on the data."""
                                            response = llm_qa.invoke(qa_prompt)
                                            answer = response.content
                                    except Exception as agent_error:
                                        # Fallback to simple LLM approach
                                        print(f"Agent error: {agent_error}")
                                        llm_qa = ChatGroq(
                                            model="llama-3.3-70b-versatile",
                                            api_key="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",
                                            temperature=0.7,
                                            max_tokens=1500,
                                        )
                                        qa_prompt = f"""You are a data analyst. Answer this question about the uploaded data:
Question: {qa_question}

Data content:
{st.session_state['summarizer_content'][:2000]}

Provide a helpful and specific answer."""
                                        response = llm_qa.invoke(qa_prompt)
                                        answer = response.content
                                        st.write(response)
                                
                                else:
                                    # For non-tabular files (PDF, DOCX, etc), use document chunking
                                    llm_qa = ChatGroq(
                                        model="llama-3.3-70b-versatile",
                                        api_key="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",
                                        temperature=0.7,
                                        max_tokens=1500,
                                        timeout=None,
                                        max_retries=2,
                                    )
                                    
                                    # Smart chunking for context retrieval
                                    content = st.session_state["summarizer_content"]
                                    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                                    
                                    # Create chunks
                                    chunks = []
                                    current_chunk = ""
                                    for para in paragraphs:
                                        if len(current_chunk) + len(para) < 2000:
                                            current_chunk += para + "\n\n"
                                        else:
                                            if current_chunk:
                                                chunks.append(current_chunk)
                                            current_chunk = para + "\n\n"
                                    if current_chunk:
                                        chunks.append(current_chunk)
                                    
                                    # Use top chunks for context
                                    relevant_context = "\n\n".join(chunks[:5])
                                    
                                    # Generate comprehensive answer
                                    qa_prompt = f"""You are an expert analyst with deep knowledge of the document content. 
                                            A user has asked a question about an uploaded document. Provide a comprehensive, detailed, and summarized response.

                                            User Question: {qa_question}

                                            relevant context {relevant_context}

                                            Please provide:
                                            1. A direct and specific answer to the question based on the document
                                            2. Supporting details and examples from the document if applicable
                                            3. Any relevant context or implications
                                            4. If the answer requires information not in the document, clearly state that
                                            5. Keep the response well-organized and easy to understand
                                            6. Do not exceed 100 words
                                            7. Give only response , not supprorting details 

                                            Format your response with clear sections and bullet points where appropriate."""
                                    
                                    response = llm_qa.invoke(qa_prompt)
                                    answer = response.content
                                
                                # Store in history
                                st.session_state["tab2_qa_history"].append({
                                    "question": qa_question,
                                    "answer": answer
                                })
                                
                                # Rerun to update the display with new Q&A
                                st.rerun()
                                
                            except Exception as e:
                                st.error(f"Error processing question: {str(e)}")
                                print(f"Q&A Error: {e}")
            
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
                print(f"File Processing Error: {e}")
        else:
            st.info("👆 Please upload a file to get started")
                

